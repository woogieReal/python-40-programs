"""
수료증 생성 후 PDF로 변환하는 코드 만들기
"""

from docx import Document
from openpyxl import load_workbook
import docx
from docx.oxml.ns import qn
# from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert # 워드를 pdf로 변환하기 위한 라이브러리를 불러온다.

# 엑셀에서 값을 읽는다.
load_wb = load_workbook("11_20/12.엑셀의_정보를_불러와_수료증_자동_생성/수료증명단.xlsx")

# 읽어온 엑셀 파일에서 활성화된 시트를 불러온다.
load_ws = load_wb.active

name_list = []
birthday_list = []
ho_list = []

for i in range(1, load_ws.max_row + 1):
    name_list.append(load_ws.cell(i, 1).value)
    birthday_list.append(load_ws.cell(i, 2).value)
    ho_list.append(load_ws.cell(i, 3).value)

print(name_list)
print(birthday_list)
print(ho_list)

# 엑셀에서 읽은 이름 리스트의 길이만큼 반복한다.
for i in range(len(name_list)):
    doc = docx.Document("11_20/12.엑셀의_정보를_불러와_수료증_자동_생성/수료증양식.docx")
    style = doc.styles['Normal']
    style.font.name = '나눔고딕'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    style.font.size = docx.shared.Pt(12)

    para = doc.add_paragraph()
    run = para.add_run('\n\n')
    run = para.add_run('              제 ' + ho_list[i] + ' 호\n')
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)

    para = doc.add_paragraph()
    run = para.add_run('\n\n')
    run = para.add_run('수  료  증')
    run.font.name = '나눔고딕'
    run.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(40)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    para = doc.add_paragraph()
    run = para.add_run('\n\n')
    run = para.add_run('        성       명: ' + name_list[i] + '\n')
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)
    run = para.add_run('        생 년 월 일: ' + birthday_list[i] + '\n')
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)
    run = para.add_run('        교 육 과 정: 파이썬과 40개의 작품들\n')
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)
    run = para.add_run('        교 육 날 짜: 2021.08.05~2021.09.09\n')
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)

    para = doc.add_paragraph()
    run = para.add_run('\n\n')
    run = para.add_run('        위 사람은 파이썬과 40개의 작품들 교육과정을\n')
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)
    run = para.add_run('        이수하였으므로 이 증서를 수여 합니다.\n')
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)

    para = doc.add_paragraph()
    run = para.add_run('\n\n\n')
    run = para.add_run('2021.09.19')
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    para = doc.add_paragraph()
    run = para.add_run('\n\n\n')
    run = para.add_run('파이썬교육기관장')
    run.font.name = '나눔고딕'
    run.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 워드파일을 생성
    doc.save('11_20/12.엑셀의_정보를_불러와_수료증_자동_생성/' + name_list[i] + '.docx')
    
    # 워드파일을 읽어 pdf로 변환
    convert('11_20/12.엑셀의_정보를_불러와_수료증_자동_생성/' + name_list[i] + '.docx',
            '11_20/12.엑셀의_정보를_불러와_수료증_자동_생성/' + name_list[i] + '.pdf')
