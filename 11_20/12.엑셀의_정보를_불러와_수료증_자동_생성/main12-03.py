"""
수료증 내용을 채운 후 저장하는 코드 만들기
"""

import docx
from docx.oxml.ns import qn
# from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# 파일을 읽어온다.
doc = docx.Document('11_20/12.엑셀의_정보를_불러와_수료증_자동_생성/수료증양식.docx')

# 기본이 되는 폰트와 글씨크기를 정한다.
style = doc.styles['Normal']
style.font.name = '나눔고딕'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
style.font.size = docx.shared.Pt(12)

# 문단을 생성하고 글씨를 입력 후 폰트와 글씨크기를 정한다.
para = doc.add_paragraph() 
run = para.add_run('\n\n') 
run = para.add_run('              제 2020-0001 호\n') 
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)

# 내용을 채워 넣는다.
para = doc.add_paragraph()
run = para.add_run('\n\n')
run = para.add_run('수  료  증') 
run.font.name = '나눔고딕'
run.bold = True
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(40)
para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

para = doc.add_paragraph()
run = para.add_run('\n\n')
run = para.add_run('        성       명: 장다인\n') 
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)
run = para.add_run('        생 년 월 일: 2017.12.12\n') 
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)
run = para.add_run('        교 육 과 정: 파이썬과 40개의 작품들\n')
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20) 
run = para.add_run('        교 육 날 짜: 2021.08.05~2021.09.09\n') 
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)

para = doc.add_paragraph() 
run = para.add_run('\n\n')
run = para.add_run('        위 사람은 파이썬과 40개의 작품들 교육과정을\n') 
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)
run = para.add_run('        이수하였으므로 이 증서를 수여 합니다.\n') 
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)

para = doc.add_paragraph()
run = para.add_run('\n\n\n')
run = para.add_run('2021.09.19') 
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)
para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

para = doc.add_paragraph()
run = para.add_run('\n\n\n')
run = para.add_run('파이썬교육기관장') 
run.font.name = '나눔고딕'
run.bold = True
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)
para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# 파일로 저장한다.
doc.save('11_20/12.엑셀의_정보를_불러와_수료증_자동_생성/수료증결과.docx')
